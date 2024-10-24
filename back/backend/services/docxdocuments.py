from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
import base64
import datetime
import time
import os
from io import BytesIO
from.helper import get_date_from_str, findrole
from petrovich.main import Petrovich
from petrovich.enums import Case, Gender
 
parent_dict={'Приказ':{"rp":"приказа","head":"Приказ","func":"order"},
            "Распоряжение":{"head":"Распоряжение","func": "disposition"},
            'План-график':{"head":"План-график","func": "plan_graph"},            
            'Отчет':{"head":"Отчет","func": "report"},
            "Итог":{"head":"Итоговый отчет","func": "final_report"},
                }
acception={1:"",2:"-", 3:"+"}
plan_podpisan= ["НЕ ОТПРАВЛЕНО НА УТВЕРЖДЕНИЕ","НА УТВЕРЖДЕНИИ","ОТКЛОНЕН","УТВЕРЖДЕН"]

passport_help=  {"main":[
                        {"name":"Наименование проекта","field":"longname", "func": "addstr"},
                        {"name":"Краткое наименование проекта","field":"name", "func": "addstr"},
                        {"name":"Директор проекта","field":"users", "func": "projectperson", "role":1},
                        {"name":"Руководитель проекта","field":"users", "func": "projectperson", "role":2},
                        {"name":"Участники проекта","field":"users", "func": "projectdeps"},
                    ],
                "content":[
                        {"name":"Основание для инициации проекта","field":"grounds", "func": "addinfostrlist"},
                        {"name":"Цель (цели) проекта","field":"target", "func": "addinfostrlist"},
                        {"name":"Задачи проекта","field":"tasks", "func": "addlistdict", "sub":"name"},
                        {"name":"Результат (результаты) проекта","field":"result", "func": "addinfostrlist"},
                        {"name":"Критерии успеха проекта","field":"criteria", "func": "addinfostrlist"},
                        {"name":"Период реализации проекта","field":["datestart","dateend"], "func": "addperiod"},
                        {"name":"Риски реализации проекта","field":"risks", "func": "addinfostrlist"},
                        {"name":"Влияние результата проекта","field":"influence", "func": "addinfostrlist"},
                        {"name":"Взаимосвязи с другими проектами и программами","field":"interactions", "func": "addinfostrlist"},
                        {"name":"Дополнительная информация","field":"dop_info", "func": "addinfostrlist"},
                      
                ]

}
def taskworkbefore(data, created, before=True):
    if before:
        res= [e for e in data if datetime.datetime.strptime(e['dateend'], '%Y-%m-%d')<=created or e['status']['id']==5 ]
    else:
        res= [e for e in data if datetime.datetime.strptime(e['dateend'], '%Y-%m-%d')>created and (datetime.datetime.strptime(e['dateend'], '%Y-%m-%d')-created).days<90  and e['status']['id']!=5 ]
    
    return len(res)>0
def dotteddate(datestr):
    datestr = datestr.split("-")    
    datestr.reverse()    
    return ".".join(datestr)

def make_rows_bold(row):
        
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def set_col_widths(table, widths):
    # widths = [Inches(1), Inches(2), Inches(1.5)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    return table
# def tablefiller(table, row):
class Filler:
    def addperiod(row,data, dictH):
        row[0].text = dictH["name"]
        row[1].text = f'{get_date_from_str(data[dictH["field"][0]])} - {get_date_from_str(data[dictH["field"][1]])}'.replace('"','')
        return row
    def addlistdict(row,data, dictH):
        row[0].text = dictH["name"]
        for (indx,x) in enumerate(data[dictH["field"]]):
                if indx==0:
                    row[1].text = x[dictH["sub"]]
                else:
                    row[1].add_paragraph(x[dictH["sub"]])

        return row
    def addinfostrlist(row,data, dictH):
        row[0].text = dictH["name"]
        
        if (isinstance(data["info"][dictH["field"]], str)):
            
            row[1].text = data["info"][dictH["field"]]
        else:
            for (indx,x) in enumerate(data["info"][dictH["field"]]):
                if indx==0:
                    row[1].text = x
                else:
                    row[1].add_paragraph(x)
        return row
    
    
    def addstr(row,data, dictH):
        row[0].text = dictH["name"]
        row[1].text = data[dictH["field"]]
        return row

    def projectperson(row,data,dictH):
        row[0].text = dictH["name"]
        person = next(filter(lambda d: d.get("role")['id'] == dictH["role"], data[dictH["field"]]), None)        
        row[1].text = person['user']['fullname']
        return row

    def projectdeps(row,data,dictH):
        row[0].text = dictH["name"]
        deps = []
        for x in data[dictH["field"]]:
            deps.append(x["user"]["department"]['name'])
        
        for (indx,x) in enumerate(list(set(deps))):
            if indx==0:
                row[1].text = x
            else:
                row[1].add_paragraph(x)
        
        return row
    









class DocxCreator:
    def docChooser(data):
        return  getattr(DocxCreator, parent_dict[data.doctype.name]['func'])(data)
    def savedoc(doc):
        fn = f'created/{time.time()}.docx'
        doc.save(fn)
        return fn

 
    
    def decodetobase(fn):  

        data=open(fn,'rb').read()        
        data = base64.b64encode(data)       
        os.remove(fn)
        return data
    
    def createDocumentWithHeader():
        document=Document()
        p = document.add_paragraph("МИНФИН РОССИИ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph("ФЕДЕРАЛЬНАЯ НАЛОГОВАЯ СЛУЖБА")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph("(ФНС РОССИИ)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return document
      

    def confirmer(data):
        
        document=Document()
        h=document.add_heading('ЛИСТ ЭЛЕКТРОННОГО СОГЛАСОВАНИЯ',1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER       
        p = document.add_paragraph(f'{parent_dict[data.doctype.name]["rp"]} {"№" + data.docnumber if (data.docnumber and len(data.docnumber)>0) else""}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h= document.add_heading(data.project.name, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=5)
        table.style=document.styles['Table Grid']
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Должность'
        hdr_cells[1].text = 'ФИО'
        hdr_cells[2].text = 'Дата'
        hdr_cells[3].text = 'Согласование'
        hdr_cells[4].text = 'Примечание'
        
        for x in data.doc_user:
            row_cells = table.add_row().cells
            row_cells[0].text = str(x.user.position.name)
            row_cells[1].text = str(x.user.fullname)
            row_cells[2].text = str(x.acceptedtime.strftime("%H:%M:%S %d.%m.%Y") if x.acceptedtime else "")
            row_cells[3].text = str(acception[x.accepted])
            row_cells[4].text = str(x.comment)
       
        return DocxCreator.savedoc(document)
    
    def order(data):
        

        pet = Petrovich()
        document = DocxCreator.createDocumentWithHeader()
        p = document.add_paragraph("ПРИКАЗ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=2)
        row=table.add_row().cells
        p=row[0].add_paragraph(get_date_from_str(data.vrersion['orderdate']))
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        p=row[1].add_paragraph(f'№ {data.docnumber}')
        p.alignment=WD_ALIGN_PARAGRAPH.RIGHT               
        p = document.add_paragraph(f'Об организации проекта {data.vrersion["longname"]}' )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(f'В целях контроля реализации направлений деятельности ФНС России в сфере {data.vrersion["info"]["project_sphere"]}, приказываю:' ) 
        p = document.add_paragraph(f'Провести Проект {data.vrersion["longname"]}',  style='List Number' )        
        p = document.add_paragraph(f'Установить период проведения Проекта с {get_date_from_str(data.vrersion["datestart"])} по {get_date_from_str(data.vrersion["dateend"])}',  style='List Number' )
        user = findrole(data.vrersion["users"],1)
        p = document.add_paragraph(f'Назначить директором Проекта {pet.lastname(user["user"]["sn"], Case.GENITIVE,None)} {pet.firstname(user["user"]["givenName"].split(" ")[0], Case.GENITIVE, None)} {pet.middlename(user["user"]["givenName"].split(" ")[1], Case.GENITIVE, None)}',style='List Number' )
            
             
        user = findrole(data.vrersion["users"],2)
        coord = findrole(data.vrersion["users"],4)
        p = document.add_paragraph(f'Назначить руководителем Проекта {pet.lastname(user["user"]["sn"], Case.GENITIVE,None)} {pet.firstname(user["user"]["givenName"].split(" ")[0], Case.GENITIVE, None)} {pet.middlename(user["user"]["givenName"].split(" ")[1], Case.GENITIVE, None)}, координатором Проекта   {pet.lastname(coord["user"]["sn"], Case.GENITIVE,None)} {pet.firstname(coord["user"]["givenName"].split(" ")[0], Case.GENITIVE, None)} {pet.middlename(coord["user"]["givenName"].split(" ")[1], Case.GENITIVE, None)}',  style='List Number' )
        p = document.add_paragraph(f'Директору проекта утвердить паспорт проекта в срок не позднее {get_date_from_str(data.vrersion["info"]["passport_last_date"])}',  style='List Number' )
        p = document.add_paragraph(f'Финансовое обеспечение расходных обязательств, связанных с реализацией настоящего приказа, осуществляется {data.vrersion["info"]["financial_support"]}',  style='List Number' )
        p = document.add_paragraph(f'Контроль за исполнением настоящего приказа возложить на {pet.lastname(data.vrersion["info"]["controller"]["sn"], Case.GENITIVE,None)} {pet.firstname(data.vrersion["info"]["controller"]["givenName"].split(" ")[0], Case.GENITIVE, None)} {pet.middlename(data.vrersion["info"]["controller"]["givenName"].split(" ")[1], Case.GENITIVE, None)}',  style='List Number' )
        footer_section = document.sections[0]
        footer = footer_section.footer
       
        if data.approver:
            short = data.approver.givenName.split(" ")
            p = document.add_paragraph(f'{data.approver.position.name}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = document.add_paragraph(f'{short[0][0]}.{short[1][0]}. {data.approver.sn}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        return DocxCreator.savedoc(document)

    def disposition(data):
        document = DocxCreator.createDocumentWithHeader()
        p = document.add_paragraph("РАСПОРЯЖЕНИЕ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=2)
        row=table.add_row().cells
        p=row[0].add_paragraph(get_date_from_str(data.vrersion['orderdate']))
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        p=row[1].add_paragraph(f'№ {data.docnumber}')
        p.alignment=WD_ALIGN_PARAGRAPH.RIGHT       
        
        p = document.add_paragraph(f'Об утверждении Паспорта Проекта, плана контрольных событий Проекта {data.vrersion["longname"]} и состава рабочей группы по реализации Проекта {data.vrersion["longname"]}' )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(f'В целях контроля реализации направлений деятельности ФНС России в сфере {data.vrersion["info"]["project_sphere"]}:' ) 
        p = document.add_paragraph(f'Утвердить:',  style='List Number' ) 
        p = document.add_paragraph(f'Паспорт Проекта {data.vrersion["longname"]} согласно приложению № 1 к настоящему распоряжению;' ) 
        p = document.add_paragraph(f'План контрольных событий Проекта {data.vrersion["longname"]} согласно приложению № 2 к настоящему распоряжению;' ) 
        p = document.add_paragraph(f'состав Рабочей группы по реализации Проекта {data.vrersion["longname"]} согласно приложению № 3 к настоящему распоряжению;' ) 
        p = document.add_paragraph(f'Контроль за исполнением настоящего распоряжения оставляю за собой',  style='List Number' ) 
        if data.approver:
            table = document.add_table(rows=1, cols=2)
            row=table.add_row().cells
            p=row[0].add_paragraph(f'{data.approver.position.name}')
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            short = data.approver.givenName.split(" ")
            p=row[1].add_paragraph(f'{short[0][0]}.{short[1][0]}. {data.approver.sn}')
            p.alignment=WD_ALIGN_PARAGRAPH.RIGHT     
            
           
        document.add_page_break()
        p = document.add_paragraph('ПРИЛОЖЕНИЕ № 1' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f'{plan_podpisan[0 if data.approved is None else data.approved]}' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph("распоряжением ФНС России" )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f"от {get_date_from_str(data.vrersion['orderdate'])}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f'№ {data.docnumber}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph("ПАСПОРТ ПРОЕКТА")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph("ОСНОВНЫЕ ПОЛОЖЕНИЯ")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table = document.add_table(0,2)
        table.style=document.styles['Table Grid']
        for (indx,x) in enumerate(passport_help["main"]):
            row = table.add_row().cells
            row = getattr(Filler, x['func'])(row,data.vrersion,x)
           
        p = document.add_paragraph("СОДЕРЖАНИЕ ПРОЕКТА")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table = document.add_table(0,2)
        table.style=document.styles['Table Grid']
        for (indx,x) in enumerate(passport_help["content"]):
            row = table.add_row().cells
            row = getattr(Filler, x['func'])(row,data.vrersion,x)
        
        p = document.add_page_break()
        
        p = document.add_paragraph('ПРИЛОЖЕНИЕ № 2' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f'{plan_podpisan[0 if data.approved is None else data.approved]}' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph("распоряжением ФНС России" )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f"от {get_date_from_str(data.vrersion['orderdate'])}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f'№ {data.docnumber}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph("План контрольных событий проекта")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(data.vrersion['longname'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(2,4)
        table.style=document.styles['Table Grid']
        row = table.rows[0].cells
        row2 = table.rows[1].cells
        
   
        for (index,x) in enumerate(["№","Наименование задачи/контрольного события","Ответственный исполнитель","Дата решения/наступления контрольного события"]):
            
            row[index].text = x
            row2[index].text = str(index+1)
        
        num = 1
        
        for (index,x) in enumerate(data.vrersion["tasks"]):
            row = table.add_row().cells
            row[0].text = str(num)
            row[1].text = x["name"]
            row[2].text = x["responsible"]["sn"]
            row[3].text = get_date_from_str(x["dateend"]).replace('"','')
            num+=1
            make_rows_bold(row)
            if len(x['task_work'])>0:
                for (index,w) in enumerate(x['task_work']):
                    row = table.add_row().cells
                    row[0].text = str(num)
                    row[1].text = w["name"]
                    row[2].text = w["responsible"]["sn"]
                    row[3].text = get_date_from_str(w["dateend"]).replace('"','')
                    num+=1


        table=set_col_widths(table,  [Inches(.1), Inches(3), Inches(1.5), Inches(1.5)])
        
        p = document.add_page_break()
        
        p = document.add_paragraph('ПРИЛОЖЕНИЕ № 3' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f'{plan_podpisan[0 if data.approved is None else data.approved]}' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph("распоряжением ФНС России" )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f"от {get_date_from_str(data.vrersion['orderdate'])}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph(f'№ {data.docnumber}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = document.add_paragraph("Состав рабочей группы по реализации проекта")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(data.vrersion['longname'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(1,5)
        table.style=document.styles['Table Grid']
        row = table.rows[0].cells
        for (index,x) in enumerate(["№","Наименование проектной роли","Процент загрузки на проекте","ФИО должностного лица","Наименование подразделения и должности"]):            
            row[index].text = x
        
        roleid=1
        sortedusers = sorted(data.vrersion['users'], key=lambda d: d['role']['id']) 
        # for x in enumerate(data.vrersion['users']):
        # print(sortedusers)
        for (index,x) in enumerate(sortedusers):
            row = table.add_row().cells
            row[0].text = str(index+1)
            row[1].text = x["role"]["name"]
            row[2].text = str(x["percent"])
            row[3].text = x["user"]["fullname"]
            row[4].text = x["user"]["department"]["name"]
            row[4].add_paragraph(x["user"]["otdel"]["name"])

        table=set_col_widths(table,  [Inches(.1), Inches(.5), Inches(.5), Inches(1.5), Inches(2.5)])







        return DocxCreator.savedoc(document)

        
    def plan_graph(data):
              
        document=Document()  

        section = document.sections
        for s in section:
            s.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = s.page_height, s.page_width
            s.page_width = new_width
            s.page_height = new_height
        p = document.add_paragraph("ПЛАН-ГРАФИК ПРОЕКТА")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = document.add_paragraph(f'{plan_podpisan[0 if data.approved is None else data.approved]}' )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if data.approver:
            p =document.add_paragraph(data.vrersion["approver"]['position']["name"] )
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            short = data.vrersion["approver"]["givenName"].split(" ")
            p = document.add_paragraph(f'{short[0][0]}.{short[1][0]}. {data.vrersion["approver"]["sn"]}')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p =document.add_paragraph(get_date_from_str(str(data.approvedtime).split(" ")[0]))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
       
        table = document.add_table(4, 9)
        table.style=document.styles['Table Grid']
        hdr1 = table.rows[2].cells
        uph =  table.rows[1].cells
        hdr2 = table.rows[3].cells
        header = ['№ п/п',"Наименование задачи/контрольного события", "Ответственный исполнитьель","Ожидаемый результат реализации мероприятия",
        "Срок начала реализации", "Срок окончания реализации (дата контрольного события)","единица измерения","значение", "Примечание" ]
        table.cell(0, 0).text = data.vrersion['longname']
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0, 0).merge(table.cell(0, 8))
        for n in range(9):
            hdr1[n].text = header[n]
            hdr2[n].text = str(n+1) if n!=0 else ""
            if n not in [6,7]:
                uph[n].merge(hdr1[n])
            elif n==6:
                uph[6].text = 'Объем ресурсного обеспечения'
                uph[6].merge(uph[7])
        
        num=1
        
        for x in data.vrersion['tasks']:
            secnum=1
            row = table.add_row().cells
            row[0].text = str(num)
            row[1].text = str(x["name"])            
            short = x["responsible"]["givenName"].split(" ")      
           
            row[2].text = f'{short[0][0]}.{short[1][0]}. {x["responsible"]["sn"]}'
            row[3].text = x['expected_result'] if 'expected_result' in x.keys() else ''
            row[4].text = dotteddate(x["datestart"])
            row[5].text = dotteddate(x["dateend"])
            row[6].text = x["unit_measurement"] if 'unit_measurement' in x.keys() else ''
            row[7].text = x["meaning"] if 'meaning' in x.keys() else ''
            row[8].text = x["remark"] if 'remark' in x.keys() else ''
            make_rows_bold(row)
            for y in x["task_work"]:
                row = table.add_row().cells
                row[0].text = f"{str(num)}.{str(secnum)}"
                row[1].text = str(y["name"])            
                short = y["responsible"]["givenName"].split(" ")      
            
                row[2].text = f'{short[0][0]}.{short[1][0]}. {y["responsible"]["sn"]}'
                row[3].text = y['expected_result'] if 'expected_result' in y.keys() else ''
                row[4].text = dotteddate(y["datestart"])
                row[5].text = dotteddate(y["dateend"])
                row[6].text = y["unit_measurement"] if 'unit_measurement' in y.keys() else ''
                row[7].text = y["meaning"] if 'meaning' in y.keys() else ''
                row[8].text = y["remark"] if 'remark' in y.keys() else ''
                secnum +=1
            num+=1


        print(data)
        return DocxCreator.savedoc(document)


    def report(data):
        print(data)  
        print("WORD")    
        document=Document()  

        section = document.sections
        for s in section:
            s.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = s.page_height, s.page_width
            s.page_width = new_width
            s.page_height = new_height
        p = document.add_paragraph("ФНС России")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(data.vrersion["longname"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph("Отчет о статусе проекта")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('на '+get_date_from_str(data.vrersion["orderdate"]).replace('"',''))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph("Исполнение контрольных событий")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        

        th=["Статус исполнения контрольных событий", "Прогноз исполнения контрольных событий на ближайшие 3 месяца"]
        tp = ["Факт/прогноз", "Прогнозный срок"]
        for (ind,el) in enumerate(data.vrersion["tasks"]):
            if len(el)>0:
                p = document.add_paragraph()
                p = document.add_paragraph(th[ind])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                table = document.add_table(1,6)
                table.style=document.styles['Table Grid']
                row = table.rows[0].cells
                for (index,x) in enumerate(["№","Блок/Контрольное событие","Плановый срок",tp[ind],"Ответственный","Комментарий"]):            
                    row[index].text = x
                
                
                
                for (index,x) in enumerate(el):
                    row = table.add_row().cells                
                    row[0].text = str(x["keyy"])
                    row[1].text = x["name"]
                    row[2].text = dotteddate(x["dateend"])
                    row[3].text =dotteddate(x["planning"])
                    short = x["responsible"]["givenName"].split(" ")    
                    row[4].text = f'{short[0][0]}.{short[1][0]}. {x["responsible"]["sn"]}'
                    row[5].text = x["comment"]["text"]
                    
                    if 'task_work' in x.keys():
                        make_rows_bold(row)
                table=set_col_widths(table,  [Inches(.1), Inches(1), Inches(1), Inches(1.5), Inches(2.5)])     
                        
        parag = ["Прочие проблемы и риски", "Ключевые результаты, не вошедшие в план"]
        th = [["№","Наименование проблемы/риска","комментарий"],["№","Ключевые достигнутые результаты","Ключевые запланированные результаты"]]
        keys = ["planresult","problemsrisks"]

        for (index,x) in enumerate(keys):
            if len(data.vrersion[x])>0:
                p = document.add_paragraph()
                p = document.add_paragraph(parag[index])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                table = document.add_table(1,3)
                table.style=document.styles['Table Grid']
                row = table.rows[0].cells
                for (i,val) in enumerate(th[index]):            
                    row[i].text = val
                for(i,val) in enumerate(data.vrersion[x]):
                    row = table.add_row().cells
                    row[0].text = str(i+1)
                    row[1].text = val["problem"]
                    row[2].text = val["comment"]
                table=set_col_widths(table,  [Inches(.1), Inches(5), Inches(5)])
        


        return DocxCreator.savedoc(document)


    def final_report(data):
        print(data.vrersion)
        document=Document() 
        p = document.add_paragraph("ИТОГОВЫЙ ОТЧЕТ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph("по проекту/этапу проекта")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        
        table = document.add_table(3,2)
        table.style=document.styles['Table Grid']
        row = table.rows[0].cells
        left=["Наименование проекта/этапа проекта","Директор проекта","Руководитель проекта"]
        right=[data.vrersion["longname"], next(filter(lambda d: d.get("role")['id'] == 1, data.vrersion["users"]), None)["user"]["fullname"],
                                          next(filter(lambda d: d.get("role")['id'] == 2, data.vrersion["users"]), None)["user"]["fullname"]]
        for i in range(3):
            row = table.rows[i].cells
            row[0].text = left[i]
            row[1].text = right[i]
            
        tables=[{"tname":'Достижение результатов проекта/этапа проекта', "t_head": ["№","Запланированные результаты","Окончательный статус","Комментарий"],  "field":"result", "bottom":"Итого достигнуто результатов"  },
                {"tname":'Достижение качества результатов проекта/этапа проекта', "t_head": ["№","Запланированные результаты","Окончательный статус","Комментарий"],  "field":"criteria", "bottom":"Итого достигнуто результатов"  }
                ]  
        for a in tables:  
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(a["tname"])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            table = document.add_table(1,4)
            table.style=document.styles['Table Grid']
            row = table.rows[0].cells
            for (index,x) in enumerate(a["t_head"]):            
                        row[index].text = x
            
            for (index,x) in enumerate(data.vrersion[a['field']]):  
                row = table.add_row().cells          
                row[0].text = str(x['index'])
                row[1].text = x['name']
                row[2].text = str(x["status"]) if "status" in x else""
                row[3].text = str(x["comment"])  if "comment" in x else""
            row = table.add_row().cells

            row[0].text = "Итого достигнуто результатов"
            row[0].merge(row[1]) 
            row[2].text = f"{data.vrersion['result_description'][a['field']]['count']}/({data.vrersion['result_description'][a['field']]['percent']}%)"
            row[3].text = str(data.vrersion['result_description'][a['field']]["comment"])
            p = document.add_paragraph()
            table=set_col_widths(table,  [Inches(.1),Inches(5), Inches(5), Inches(5)])
        


        
        p = document.add_paragraph()        
        p = document.add_paragraph("Соблюдение сроков проекта/этапа проекта")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        table = document.add_table(1,3)
        table.style=document.styles['Table Grid']
        row = table.rows[0].cells
        for (index,x) in enumerate(['Плановая длительность проекта/этапа проета', "Фактическая длительность проекта/этапа проекта", "Отклонение"]):            
                    row[index].text = x
        datestart=get_date_from_str(data.vrersion['datestart']).replace('"','')
        dateend = get_date_from_str(data.vrersion['dateend']).replace('"','')
        fdatestart=get_date_from_str(data.vrersion['srok']['fact']['datestart']).replace('"','')
        fdateend = get_date_from_str(data.vrersion['srok']['fact']['dateend']).replace('"','')
        diffd = (datetime.datetime.strptime(data.vrersion['dateend'], '%Y-%m-%d')-datetime.datetime.strptime(data.vrersion['datestart'], '%Y-%m-%d')).days
        fdiffd = (datetime.datetime.strptime(data.vrersion['srok']['fact']['dateend'], '%Y-%m-%d')-datetime.datetime.strptime(data.vrersion['srok']['fact']['datestart'], '%Y-%m-%d')).days
        deviation = round((fdiffd-diffd)*100/diffd, 2)
        row = table.add_row().cells 
        row[0].text =f"{datestart}-{dateend} ({diffd}д.)"
        row[1].text= f"{fdatestart}-{fdateend} ({fdiffd}д.)"
        row[2].text = f"{deviation}%"
        table = document.add_table(2,1)
        table.style=document.styles['Table Grid']
        row = table.rows[0].cells
        row[0].text ="Дополнительные комментарии к соблюдению сроков завершения проекта"
        row = table.rows[1].cells
        row[0].text =data.vrersion['srok']['comment']
        return DocxCreator.savedoc(document)


